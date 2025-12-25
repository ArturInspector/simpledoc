"""
Генератор DOCX документов из шаблонов
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any
import os
import re
from jinja2 import Template


def replace_text_in_document(doc: Document, replacements: Dict[str, str]):
    """
    Заменяет текст в документе по словарю замен
    
    Args:
        doc: Document объект
        replacements: Словарь {ключ: значение} для замены
    """
    # Замена в параграфах
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
    
    # Замена в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))


def generate_docx(
    template_path: str,
    output_path: str,
    data: Dict[str, Any]
) -> str:
    """
    Генерирует DOCX документ из шаблона с подстановкой данных
    
    Args:
        template_path: Путь к шаблону .docx
        output_path: Путь для сохранения результата
        data: Словарь с данными для подстановки
        
    Returns:
        Путь к сгенерированному файлу
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
    
    # Загружаем шаблон
    doc = Document(template_path)
    
    # Формируем словарь замен в формате {{key}}
    replacements = {}
    for key, value in data.items():
        replacements[f"{{{{{key}}}}}}"] = str(value)
        replacements[f"{{{{ {key} }}}}"] = str(value)  # С пробелами
    
    # Выполняем замены
    replace_text_in_document(doc, replacements)
    
    # Создаем директорию если нужно
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Сохраняем документ
    doc.save(output_path)
    return output_path


def generate_docx_from_template(
    template_content: str,
    output_path: str,
    data: Dict[str, Any]
) -> str:
    """
    Генерирует DOCX из текстового шаблона (более продвинутый вариант)
    
    Args:
        template_content: Содержимое шаблона в формате Jinja2
        output_path: Путь для сохранения
        data: Данные для подстановки
        
    Returns:
        Путь к сгенерированному файлу
    """
    # Рендерим шаблон через Jinja2
    template = Template(template_content)
    rendered_content = template.render(**data)
    
    # Создаем новый документ
    doc = Document()
    
    # Парсим отрендеренный контент и добавляем в документ
    # Это упрощенная версия, можно улучшить для поддержки форматирования
    for line in rendered_content.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph()
    
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path

